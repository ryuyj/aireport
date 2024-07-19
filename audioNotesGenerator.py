from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import os


load_dotenv()

class AudioNotesGenerator:
    def __init__(self):
        self.client = OpenAI(api_key=os.environ.get("API_KEY"))

    def transcribe_audio(self, audio_file_path):
        with open(audio_file_path, 'rb') as audio_file: # 오디오 파일 열기
            transcription = self.client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
                ) # Whisper 모델로 트랜스크립션 요청
        return self.remove_str(transcription.text) # 트랜스크립션 결과 반환

    def audio_notes(self, transcription):
        abstract_summary = self.abstract_summary_extraction(self.remove_str(transcription))
        key_points = self.key_points_extraction(self.remove_str(transcription)) # 주요 요점 추출
        action_items = self.action_item_extraction(self.remove_str(transcription)) # 작업 항목 추출
        sentiment = self.sentiment_analysis(self.remove_str(transcription)) # 감정 분석
        return { 
            'abstract_summary': abstract_summary, # 요약 반환
            'key_points': key_points, # 주요 요점 반환
            'action_items': action_items, # 작업 항목 반환
            'sentiment': sentiment # 감정 분석 반환
        }
    
    def remove_str(self, s):
        # 문자열의 길이를 구합니다.
        length = len(s)
        # 마지막 10%의 길이를 계산합니다.
        cut_length = length // 50
        # 문자열에서 마지막 10%를 자릅니다.
        return s[:-cut_length]

    def abstract_summary_extraction(self, transcription):
        response = self.client.chat.completions.create(
            model="gpt-4", # GPT-4 모델 사용
            temperature=0, # 출력을 제어하는 매개변수
            messages=[
                {
                    "role": "system",
                    "content": "You are a highly skilled AI trained in language comprehension and summarization..." 
                },
                {
                    "role": "user",
                    "content": transcription # 사용자 메시지로 트랜스크립션 전달
                }
            ]
            )
        return response.choices[0].message.content # 요약 결과 반환

    def key_points_extraction(self, transcription):
        response = self.client.chat.completions.create(
            model="gpt-4", # GPT-4 모델 사용
            temperature=0, # 출력을 제어하는 매개변수
            messages=[
                {
                    "role": "system",
                    "content": "You are a proficient AI with a specialty in distilling information into key points..."
                },
                {
                    "role": "user",
                    "content": transcription # 사용자 메시지로 트랜스크립션 전달
                }
            ]
        )
        return response.choices[0].message.content # 주요 요점 결과 반환

    def action_item_extraction(self, transcription):
        response = self.client.chat.completions.create(
            model="gpt-4", # GPT-4 모델 사용
            temperature=0, # 출력을 제어하는 매개변수
            messages=[
                {
                    "role": "system",
                    "content": "You are an AI expert in analyzing conversations and extracting action items..."
                },
                {
                    "role": "user",
                    "content": transcription # 사용자 메시지로 트랜스크립션 전달
                }
            ]
        )
        return response.choices[0].message.content # 작업 항목 결과 반환

    def sentiment_analysis(self, transcription):

        response = self.client.chat.completions.create(
            model="gpt-4", # GPT-4 모델 사용
            temperature=0, # 출력을 제어하는 매개변수
            messages=[
                {
                    "role": "system",
                    "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment..."
                },
                {
                    "role": "user",
                    "content": transcription # 사용자 메시지로 트랜스크립션 전달
                }
            ]
        )
        return response.choices[0].message.content # 감정 분석 결과 반환

    def save_as_docx(self, notes, filename):

        doc = Document() # 새 Word 문서 생성
        
        for key, value in notes.items(): # 메모의 각 항목에 대해
		    # 제목 생성
            heading = ' '.join(word.capitalize() for word in key.split('_'))
            doc.add_heading(heading, level=1) # 제목 추가
            doc.add_paragraph(value) # 내용 추가
            doc.add_paragraph() # 섹션 간에 줄 바꿈 추가
        doc.save(filename) # 문서 저장

    def generate_audio_notes(self, audio_file_path, output_filename):
        transcription = self.transcribe_audio(audio_file_path) # 오디오 파일 트랜스크립션
        notes = self.audio_notes(transcription) # 트랜스크립션으로 메모 생성
        self.save_as_docx(notes, output_filename) # 메모를 Word 문서로 저장
        return notes

if __name__ == "__main__":

    # 오디오 파일 경로 설정
    #file_path = "memo1.m4a"
    audio_file_path = "test.m4a"
    # 출력 파일 이름 설정
    output_filename = "notes.docx"
    # AudioNotesGenerator 인스턴스 생성
    generator = AudioNotesGenerator()
    # 메모 생성 및 저장
    generator.generate_audio_notes(audio_file_path, output_filename)
